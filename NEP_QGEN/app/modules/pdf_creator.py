from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
from typing import Dict, List, Any


class PDFCreator:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()

    def _setup_custom_styles(self):
        """Setup custom styles for the document"""
        # Title style
        self.title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )

        # Question style
        self.question_style = ParagraphStyle(
            'Question',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            spaceBefore=8,
            fontName='Helvetica'
        )

        # Answer style
        self.answer_style = ParagraphStyle(
            'Answer',
            parent=self.styles['Normal'],
            fontSize=10,
            spaceAfter=8,
            leftIndent=20,
            fontName='Helvetica',
            textColor=colors.blue
        )

        # Header style
        self.header_style = ParagraphStyle(
            'Header',
            parent=self.styles['Normal'],
            fontSize=12,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )

    def create_pdf(self, paper: Any, questions: List[Dict[str, Any]]) -> str:
        """Create PDF question paper"""

        # Create filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"paper_{paper.id}_{timestamp}.pdf"
        file_path = os.path.join("generated_papers", filename)

        # Ensure directory exists
        os.makedirs("generated_papers", exist_ok=True)

        # Create PDF document
        doc = SimpleDocTemplate(file_path, pagesize=A4, topMargin=1 * inch)

        # Build content
        content = []

        # Header
        content.extend(self._create_header(paper))

        # Instructions
        content.extend(self._create_instructions(paper, questions))

        # Questions
        content.extend(self._create_questions_section(questions))

        # Answer key (optional - can be toggled)
        # content.extend(self._create_answer_key(questions))

        # Build PDF
        doc.build(content)

        return file_path

    def create_docx(self, paper: Any, questions: List[Dict[str, Any]]) -> str:
        """Create DOCX question paper"""

        # Create filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"paper_{paper.id}_{timestamp}.docx"
        file_path = os.path.join("generated_papers", filename)

        # Create document
        doc = Document()

        # Header
        self._add_docx_header(doc, paper)

        # Instructions
        self._add_docx_instructions(doc, paper, questions)

        # Questions
        self._add_docx_questions(doc, questions)

        # Save document
        doc.save(file_path)

        return file_path

    def _create_header(self, paper: Any) -> List:
        """Create PDF header section"""
        content = []

        # University/Institution name (customizable)
        content.append(Paragraph("XYZ UNIVERSITY", self.title_style))
        content.append(Spacer(1, 12))

        # Paper details table
        paper_data = [
            ["Subject:", paper.subject, "Date:", datetime.now().strftime("%d/%m/%Y")],
            ["Paper Title:", paper.title, "Time:", "3 Hours"],
            ["Max Marks:", str(paper.total_marks), "Difficulty:", f"{paper.difficulty_level}/10"]
        ]

        table = Table(paper_data, colWidths=[1 * inch, 2 * inch, 1 * inch, 1.5 * inch])
        table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
        ]))

        content.append(table)
        content.append(Spacer(1, 20))

        return content

    def _create_instructions(self, paper: Any, questions: List[Dict[str, Any]]) -> List:
        """Create instructions section"""
        content = []

        # Instructions title
        content.append(Paragraph("<b>INSTRUCTIONS:</b>", self.question_style))

        # Generate instructions based on question types
        instructions = [
            "1. Read all questions carefully before attempting.",
            "2. All questions are compulsory.",
            "3. Write your answers clearly and legibly.",
            "4. Time management is crucial for completion.",
            f"5. This paper contains {len(questions)} questions for {paper.total_marks} marks."
        ]

        # Add specific instructions based on question types
        question_types = set(q.get("type", "") for q in questions)

        if "mcq" in question_types:
            instructions.append("6. For Multiple Choice Questions, select the best answer.")

        if any(q.get("marks", 0) >= 10 for q in questions):
            instructions.append("7. For long answer questions, provide detailed explanations with examples.")

        instructions.append(
            "8. Follow NEP 2020 guidelines: Focus on understanding, application, and critical thinking.")

        for instruction in instructions:
            content.append(Paragraph(instruction, self.styles['Normal']))

        content.append(Spacer(1, 20))
        content.append(Paragraph("_" * 80, self.styles['Normal']))
        content.append(Spacer(1, 20))

        return content

    def _create_questions_section(self, questions: List[Dict[str, Any]]) -> List:
        """Create questions section"""
        content = []

        # Group questions by type and marks
        grouped_questions = self._group_questions(questions)

        question_number = 1

        for group_name, group_questions in grouped_questions.items():
            # Section header
            content.append(
                Paragraph(f"<b>SECTION {chr(64 + len(content) // 10 + 1)}: {group_name}</b>", self.question_style))
            content.append(Spacer(1, 10))

            for question in group_questions:
                # Question text with number and marks
                question_text = f"Q{question_number}. {question['text']} [{question['marks']} marks]"
                content.append(Paragraph(question_text, self.question_style))

                # Add options for MCQ
                if question.get("type") == "mcq" and question.get("options"):
                    for i, option in enumerate(question["options"][:4]):  # Limit to 4 options
                        option_text = f"({chr(97 + i)}) {option}"
                        content.append(Paragraph(option_text, self.styles['Normal']))

                # Add space for answers
                if question.get("marks", 0) <= 2:
                    content.append(Spacer(1, 30))
                elif question.get("marks", 0) <= 5:
                    content.append(Spacer(1, 60))
                else:
                    content.append(Spacer(1, 100))

                question_number += 1

            content.append(Spacer(1, 20))

        return content

    def _group_questions(self, questions: List[Dict[str, Any]]) -> Dict[str, List]:
        """Group questions by type and marks"""
        groups = {}

        for question in questions:
            q_type = question.get("type", "general")
            marks = question.get("marks", 0)

            if q_type == "mcq":
                group_name = "Multiple Choice Questions"
            elif marks <= 2:
                group_name = "Short Answer Questions"
            elif marks <= 5:
                group_name = "Medium Answer Questions"
            else:
                group_name = "Long Answer Questions"

            if group_name not in groups:
                groups[group_name] = []
            groups[group_name].append(question)

        return groups

    def _create_answer_key(self, questions: List[Dict[str, Any]]) -> List:
        """Create answer key section (optional)"""
        content = []

        content.append(PageBreak())
        content.append(Paragraph("<b>ANSWER KEY (For Reference)</b>", self.title_style))
        content.append(Spacer(1, 20))

        for i, question in enumerate(questions, 1):
            answer_text = f"Q{i}. {question.get('answer', 'Answer not provided')}"
            content.append(Paragraph(answer_text, self.answer_style))
            content.append(Spacer(1, 10))

        return content

    def _add_docx_header(self, doc: Document, paper: Any):
        """Add header to DOCX document"""
        # Title
        title = doc.add_heading('XYZ UNIVERSITY', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Paper details
        doc.add_paragraph(f"Subject: {paper.subject}")
        doc.add_paragraph(f"Paper Title: {paper.title}")
        doc.add_paragraph(f"Max Marks: {paper.total_marks}")
        doc.add_paragraph(f"Time: 3 Hours")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y')}")
        doc.add_paragraph("_" * 50)

    def _add_docx_instructions(self, doc: Document, paper: Any, questions: List[Dict[str, Any]]):
        """Add instructions to DOCX document"""
        doc.add_heading('INSTRUCTIONS:', level=2)

        instructions = [
            "1. Read all questions carefully before attempting.",
            "2. All questions are compulsory.",
            "3. Write your answers clearly and legibly.",
            "4. Time management is crucial for completion.",
            f"5. This paper contains {len(questions)} questions for {paper.total_marks} marks.",
            "6. Follow NEP 2020 guidelines: Focus on understanding, application, and critical thinking."
        ]

        for instruction in instructions:
            doc.add_paragraph(instruction)

        doc.add_paragraph("_" * 50)

    def _add_docx_questions(self, doc: Document, questions: List[Dict[str, Any]]):
        """Add questions to DOCX document"""
        grouped_questions = self._group_questions(questions)
        question_number = 1

        for group_name, group_questions in grouped_questions.items():
            doc.add_heading(f"SECTION: {group_name}", level=2)

            for question in group_questions:
                # Question
                question_para = doc.add_paragraph(f"Q{question_number}. {question['text']} [{question['marks']} marks]")
                question_para.style = 'Normal'

                # MCQ options
                if question.get("type") == "mcq" and question.get("options"):
                    for i, option in enumerate(question["options"][:4]):
                        doc.add_paragraph(f"({chr(97 + i)}) {option}")

                # Space for answer
                doc.add_paragraph("\n" * (question.get("marks", 1)))

                question_number += 1
